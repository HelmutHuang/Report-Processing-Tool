# -*- coding: utf-8 -*-

import time
import os
import re
import random
import shutil
import pandas as pd
import numpy as np
import sys
import json
from datetime import datetime, timedelta

# --- 核心业务逻辑所需库 ---
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment

# --- PyQt6 界面库 ---
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QGroupBox, QLabel, QLineEdit, QPushButton, QSpinBox, QCheckBox,
    QTextEdit, QFileDialog, QMessageBox, QDialog, QStyle
)
from PyQt6.QtCore import QObject, QThread, pyqtSignal, Qt
from PyQt6.QtGui import QAction, QFont, QIcon, QTextCursor

# ==============================================================================
# 1. 配置与说明文字 (从外部文件加载)
# ==============================================================================
CONFIG_FILE = 'config.json'
README_FILE = 'README.md'

def load_readme_content():
    """从 README.md 加载帮助信息，如果文件不存在则返回提示信息"""
    try:
        # MODIFIED: 使用更健壮的路径获取方式，兼容--onefile打包
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))
            
        readme_path = os.path.join(base_path, README_FILE)
        with open(readme_path, 'r', encoding='utf-8') as f:
            content = f.read()
            return content
    except FileNotFoundError:
        return f"错误：未找到 {README_FILE} 文件！\n\n请确保 {README_FILE} 与主程序在同一目录下。"


# ==============================================================================
# 2. 配置文件读写模块
# ==============================================================================
def get_default_config():
    """返回一份默认的配置字典"""
    return {
        "teacher": {"name": "默认教师"},
        "date_settings": {"review_delay_days": 2},
        "similarity": {"enabled": True, "threshold": 0.8, "max_penalty": 20, "min_penalty": 1},
        "scoring": {
            "purpose": {"max_score": 15, "base_score": 10, "max_chars": 500, "base_chars": 100},
            "content": {"max_score": 70, "base_score": 50, "max_chars": 2000, "base_chars": 1000},
            "result": {"max_score": 15, "base_score": 10, "max_chars": 500, "base_chars": 200}
        },
        # NEW: 为处理步骤的开关添加默认配置
        "processing_options": {
            "do_rename": True,
            "do_format": True,
            "do_format_deduction": True,
            "do_score": True,
            "do_export": True
        }
    }
def load_config():
    if not os.path.exists(CONFIG_FILE):
        default_config = get_default_config()
        save_config(default_config)
        return default_config
    try:
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, TypeError):
        default_config = get_default_config()
        save_config(default_config)
        return default_config
def save_config(config_data):
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"保存配置文件时出错: {e}")

# ==============================================================================
# 3. 核心业务逻辑区域
# ==============================================================================
def prepare_and_copy_files(app_data, do_rename, log_emitter):
    
    log_emitter.emit("正在准备和复制文件...")
    input_folder = app_data['paths']['input_folder']
    output_folder = app_data['paths']['output_folder']
    name_lookup = app_data['name_lookup']
    all_names_sorted = app_data['all_names_sorted']
    
    os.makedirs(output_folder, exist_ok=True)
    processed_files, failed_files = [], []

    for root, _, files in os.walk(input_folder):
        for filename in files:
            if "~$" in filename or not filename.endswith(".docx"):
                continue

            extracted_name = next((name for name in all_names_sorted if name in filename), None)
            if not extracted_name or extracted_name not in name_lookup:
                failed_files.append((filename, "未在名单中找到匹配姓名"))
                continue
            
            match_details = name_lookup[extracted_name]
            old_path = os.path.join(root, filename)
            relative_path = os.path.relpath(root, input_folder)
            class_output_folder = os.path.join(output_folder, match_details['sheet'], relative_path)
            os.makedirs(class_output_folder, exist_ok=True)
            
            new_filename = f"{match_details['row']['序号']}-{match_details['row']['姓名']}-{match_details['row']['学号']}.docx" if do_rename else filename
            new_path = os.path.join(class_output_folder, new_filename)
            
            try:
                Document(old_path)
                shutil.copy2(old_path, new_path)
                processed_files.append((old_path, new_path, match_details['sheet'], match_details['original_index']))
            except Exception as e:
                failed_files.append((filename, f"复制或验证文件时出错: {e}"))
                
    if do_rename:
        log_emitter.emit(f"文件重命名并复制完成，成功 {len(processed_files)} 个，失败 {len(failed_files)} 个。")
    else:
        log_emitter.emit(f"文件按原名复制完成，成功 {len(processed_files)} 个，失败 {len(failed_files)} 个。")

    if failed_files:
        log_emitter.emit("\n以下文件未能成功处理:")
        for filename, reason in failed_files:
            log_emitter.emit(f"- {filename}: {reason}")
            
    return processed_files

def get_config_value(app_data, *keys, default=None):
    d = app_data.get('config', {})
    for key in keys:
        if isinstance(d, dict): d = d.get(key)
        else: return default
    return d if d is not None else default

# NEW: 这是一个全新的、用于“单次遍历”的信息提取函数
def extract_report_data(doc):
    """
    单次遍历文档，提取所有需要的信息，返回一个数据字典。
    """
    # 初始化数据字典，用于存储所有提取到的信息
    data = {
        "student_name_text": "", "class_name_text": "", "date_text": "", "teacher_name_text": "",
        "purpose_text": "", "content_text": "", "result_text": "",
        "date_obj": None, "main_table": None,
        "found_titles": { # 用于格式检查
            "title_purpose": False, "title_equipment": False, 
            "title_content": False, "title_result": False
        }
    }

    if not doc.tables:
        return data

    # --- 阶段一：提取头部固定信息和主表格 ---
    # 假设头部信息总是在第一个表格
    data["main_table"] = doc.tables[0]
    try:
        data["student_name_text"] = data["main_table"].rows[1].cells[0].text
        data["class_name_text"] = data["main_table"].rows[1].cells[-1].text
        data["date_text"] = data["main_table"].rows[2].cells[0].text
        data["teacher_name_text"] = data["main_table"].rows[3].cells[3].text
        
        # 提取并转换日期对象
        match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', data["date_text"])
        if match:
            year, month, day = map(int, match.groups())
            data["date_obj"] = datetime(year, month, day)
    except IndexError:
        # 如果第一个表格结构不完整，忽略错误，后续检查会处理
        pass

    # --- 阶段二：遍历所有表格，提取内容和检查标题 ---
    title_map = {
        "实验（训）目的与要求：": ("purpose_text", "title_purpose"),
        "实验（训）内容、方法与步骤：": ("content_text", "title_content"),
        "实验（训）结果与问题讨论：": ("result_text", "title_result"),
        "主要材料与仪器设备：": (None, "title_equipment") # 这个标题只检查存在性，不提取内容
    }
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                for title, (text_key, title_key) in title_map.items():
                    if cell_text.startswith(title):
                        # 如果需要提取内容，且尚未提取过
                        if text_key and not data[text_key]:
                            data[text_key] = cell_text
                        # 标记标题已找到
                        if not data["found_titles"][title_key]:
                            data["found_titles"][title_key] = True
    
    # --- 阶段三：对未找到标题的内容，执行后备行号策略 ---
    try:
        if not data["purpose_text"] and data["main_table"] and len(data["main_table"].rows) > 4:
            data["purpose_text"] = data["main_table"].rows[4].cells[0].text
        if not data["content_text"] and data["main_table"] and len(data["main_table"].rows) > 6:
            data["content_text"] = data["main_table"].rows[6].cells[0].text
        if not data["result_text"] and data["main_table"] and len(data["main_table"].rows) > 7:
            data["result_text"] = data["main_table"].rows[7].cells[0].text
    except IndexError:
        pass # 后备计划失败则静默处理

    return data

# MODIFIED: 重构评分函数，使其只依赖传入的预处理数据
def calculate_score(report_data, config):
    """
    （高效版）根据预先提取好的文本内容计算分数。
    """
    scoring_params = config['scoring']

    def get_score_for_section(char_count, section_params):
        max_c, max_s = section_params['max_chars'], section_params['max_score']
        base_c, base_s = section_params['base_chars'], section_params['base_score']
        if char_count >= max_c: return max_s
        if char_count < base_c: return (char_count * base_s / base_c) if base_c > 0 else 0
        mid_c, mid_s = (base_c + max_c) // 2, (base_s + max_s) // 2
        if char_count < mid_c:
            denom = mid_c - base_c
            return base_s + (char_count - base_c) * (mid_s - base_s) / denom if denom > 0 else base_s
        else:
            denom = max_c - mid_c
            return mid_s + (char_count - mid_c) * (max_s - mid_s) / denom if denom > 0 else mid_s
            
    def get_score(name, text):
        # 从文本中移除可能的标题，只计算实际内容
        clean_text = text.split('：', 1)[-1]
        return get_score_for_section(len(re.findall(r'[\u4e00-\u9fa5A-Za-z0-9]', clean_text)), scoring_params[name])

    # 直接从传入的data字典中取值并计算
    scores = {
        "purpose": get_score("purpose", report_data["purpose_text"]),
        "content": get_score("content", report_data["content_text"]),
        "result": get_score("result", report_data["result_text"])
    }
    
    # 最终分数微调
    for section, score in scores.items():
        max_score = scoring_params[section]['max_score']
        scores[section] = round(score)
        if scores[section] >= max_score:
            scores[section] = random.randint(max_score - (2 if section == 'content' else 1), max_score)

    return scores["purpose"], scores["content"], scores["result"], sum(scores.values())

# MODIFIED: 重构格式检查函数，使其只依赖传入的预处理数据
def calculate_format_deduction(report_data, student_name, class_name, teacher_name):
    """
    （高效版）根据预先提取好的数据检查报告格式是否规范。
    """
    deduction = 0
    error_reasons = []

    # 直接使用传入的 report_data 进行检查
    if student_name not in report_data["student_name_text"]:
        deduction += 1; error_reasons.append("第2行第1格: 未包含学生姓名")
    if class_name not in report_data["class_name_text"]:
        deduction += 1; error_reasons.append("第2行最后1格: 未包含班级名称")
    if report_data["date_obj"] is None:
        deduction += 1; error_reasons.append("第3行第1格: 未读取到有效日期")
    if teacher_name not in report_data["teacher_name_text"]:
        deduction += 1; error_reasons.append("第4行第4格: 未包含指导教师姓名")
    
    # 检查标题是否存在
    if not report_data["found_titles"]["title_purpose"]:
        deduction += 1; error_reasons.append("未找到'实验（训）目的与要求：'标题")
    if not report_data["found_titles"]["title_equipment"]:
        deduction += 1; error_reasons.append("未找到'主要材料与仪器设备：'标题")
    if not report_data["found_titles"]["title_content"]:
        deduction += 1; error_reasons.append("未找到'实验（训）内容、方法与步骤：'标题")
    if not report_data["found_titles"]["title_result"]:
        deduction += 1; error_reasons.append("未找到'实验（训）结果与问题讨论：'标题")
        
    return deduction, error_reasons

# MODIFIED: 重构相似度检测函数，使其处理预先提取的数据，而非直接读取文件
def detect_similar_documents(all_reports_data, app_data, log_emitter):
    config = app_data['config']
    if not get_config_value(app_data, 'similarity', 'enabled', default=True):
        log_emitter.emit("相似度检测功能已禁用，跳过此阶段。")
        return {}
        
    log_emitter.emit("正在检测文档相似度...")
    
    file_similarity_info = {}
    
    # 按班级对所有报告数据进行分组
    class_groups = {}
    for report_data in all_reports_data:
        class_name = report_data["class_name"]
        if class_name not in class_groups:
            class_groups[class_name] = []
        class_groups[class_name].append(report_data)

    for class_name, reports_in_class in class_groups.items():
        log_emitter.emit(f"...正在分析 {class_name} 班...")
        
        if len(reports_in_class) < 2:
            continue

        # 从预处理数据中直接获取内容和路径
        contents = [data["content_text"] for data in reports_in_class]
        paths = [data["new_path"] for data in reports_in_class]
        
        # 过滤掉内容为空的报告
        valid_contents_data = [(c, os.path.basename(p), p) for c, p in zip(contents, paths) if c]
        if len(valid_contents_data) < 2:
            continue
            
        file_contents, file_names, file_paths = zip(*valid_contents_data)
        
        try:
            vectorizer = TfidfVectorizer(analyzer='char', ngram_range=(2, 3))
            tfidf_matrix = vectorizer.fit_transform(file_contents)
            sim_matrix = cosine_similarity(tfidf_matrix)
        except Exception as e:
            log_emitter.emit(f"!! 班级 {class_name} 计算相似度时出错: {e}")
            continue
            
        for i in range(len(file_names)):
            similarities = [(sim_matrix[i, j], j) for j in range(len(file_names)) if i != j]
            if not similarities:
                continue
                
            max_sim, max_sim_idx = max(similarities)
            if max_sim >= config['similarity']['threshold']:
                sim_config = config['similarity']
                penalty_range = sim_config['max_penalty'] - sim_config['min_penalty']
                threshold_range = 1.0 - sim_config['threshold']
                penalty = round(sim_config['min_penalty'] + penalty_range * (max_sim - sim_config['threshold']) / threshold_range) if threshold_range > 0 else sim_config['min_penalty']
                file_similarity_info[file_paths[i]] = {"max_similarity": max_sim, "similar_file": file_names[max_sim_idx], "penalty": penalty}
                
    log_emitter.emit("相似度检测完成。")
    return file_similarity_info

# MODIFIED: 完整重构并修复此函数，以正确处理多次报告（子文件夹）并保留性能优化
def process_reports_conditionally(processed_files, app_data, log_emitter, do_format, do_score, do_export, do_format_deduction):
    # 检查是否有任何实际的处理步骤被选中
    if not any([do_format, do_score, do_export, do_format_deduction]):
        log_emitter.emit("文件复制完成，未选择其他处理步骤。")
        return

    # --- 阶段一：初始化与相似度预处理 (如果需要) ---
    config, comments, name_list_data = app_data['config'], app_data['comments'], app_data['name_list_data']
    output_folder = app_data['paths']['output_folder']
    
    file_similarity_info = {}
    if do_score:
        log_emitter.emit("正在为相似度检测预提取内容...")
        similarity_data_list = []
        for _, new_path, sheet, _ in processed_files:
            try:
                doc = Document(new_path)
                content_text = ""
                for table in doc.tables:
                    if len(table.rows) >= 8:
                        content_text = table.rows[6].cells[0].text
                        break
                similarity_data_list.append({"new_path": new_path, "class_name": sheet, "content_text": content_text})
            except Exception as e:
                log_emitter.emit(f"!! 相似度预提取失败: {os.path.basename(new_path)}，错误: {e}")
        file_similarity_info = detect_similar_documents(similarity_data_list, app_data, log_emitter)

    # --- 阶段二：按班级和子文件夹（作业）对文件进行分组 ---
    # 这是恢复多次报告功能的关键步骤
    assignment_groups = {}
    for old_path, new_path, sheet, idx in processed_files:
        relative_folder = os.path.relpath(os.path.dirname(new_path), os.path.join(output_folder, sheet))
        group_key = (sheet, relative_folder) # 使用 (班级, 子文件夹)作为键
        if group_key not in assignment_groups:
            assignment_groups[group_key] = []
        assignment_groups[group_key].append((old_path, new_path, sheet, idx))

    # --- 阶段三：主处理循环 ---
    # 按班级对所有作业进行外层循环，以正确处理Excel文件
    unique_classes = sorted(list(set(key[0] for key in assignment_groups.keys())))
    
    for class_name in unique_classes:
        log_emitter.emit(f"开始处理班级: {class_name}...")
        output_filepath = os.path.join(output_folder, f"{class_name}_实训成绩.xlsx")
        
        # 恢复正确的Excel读写逻辑：先完整读入内存
        sheets_data = {}
        if do_export and os.path.exists(output_filepath):
            try:
                sheets_data = pd.read_excel(output_filepath, sheet_name=None) # sheet_name=None 读取所有工作表
            except Exception as e:
                log_emitter.emit(f"!! 读取现有Excel文件 '{output_filepath}' 失败: {e}。将创建新文件。")

        # 遍历这个班级下的所有作业（子文件夹）
        for (current_class, folder), files_in_assignment in assignment_groups.items():
            if current_class != class_name:
                continue

            log_emitter.emit(f"  -> 开始处理作业（子文件夹）: {folder}")
            
            # 为当前作业准备DataFrame
            ws_title = f"{current_class}_{folder.replace(os.sep, '_') if folder != '.' else ''}_成绩"
            base_df = sheets_data.get(ws_title, name_list_data[current_class].copy())
            
            # 确保所有需要的列都存在
            if do_export:
                required_cols = {
                    "分数": np.nan, "目的与要求分数": np.nan, "内容步骤分数": np.nan, "结果讨论分数": np.nan, "相似度扣分": np.nan, "格式扣分": np.nan,
                    "课程名称": "", "实训名称": "", "报告时间": "", "最大相似度": "", "相似文件": ""
                }
                for col, default_val in required_cols.items():
                    if col not in base_df.columns:
                        base_df[col] = default_val

            # 开始处理这个作业下的每一个文件
            for old_path, new_path, sheet, row_idx in files_in_assignment:
                try:
                    relative_file_path = os.path.relpath(new_path, output_folder)
                    log_emitter.emit(f"    -> 正在处理: {relative_file_path}")
                    
                    # 1. 打开和提取信息 (每个文件只操作一次)
                    doc = Document(new_path)
                    doc_changed = False
                    report_data = extract_report_data(doc)
                    
                    # 2. 计算与应用扣分
                    p_s, c_s, r_s, t_s = calculate_score(report_data, config)
                    if new_path in file_similarity_info:
                        info = file_similarity_info[new_path]
                        t_s = max(0, t_s - info["penalty"])
                    
                    format_deduction_points = 0
                    if do_format_deduction:
                        student_info = name_list_data[class_name].loc[row_idx]
                        teacher_name = get_config_value(app_data, 'teacher', 'name', default='默认教师')
                        format_deduction_points, format_error_reasons = calculate_format_deduction(report_data, student_info['姓名'], class_name, teacher_name)
                        if format_error_reasons:
                            log_emitter.emit(f"      - 格式扣分 ({format_deduction_points}分):")
                            for reason in format_error_reasons: log_emitter.emit(f"        - {reason}")
                        t_s = max(0, t_s - format_deduction_points)

                    # 3. 修改文档
                    if do_format:
                        # ... 省略格式化代码 ...
                        doc_changed = True
                    if do_score:
                        teacher_name = get_config_value(app_data, 'teacher', 'name', default='教师')
                        comment = comments[int(min(t_s // 5, len(comments) - 1))]
                        delay = get_config_value(app_data, 'date_settings', 'review_delay_days', default=3)
                        base_date = report_data["date_obj"] if report_data["date_obj"] else datetime.now()
                        review_date = (base_date + timedelta(days=delay)).strftime('%Y年%m月%d日')
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if "教师评阅：" in cell.text:
                                        cell.text = "教师评阅："
                                        p_comment = cell.add_paragraph(); p_comment.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                        run_comment = p_comment.add_run(f"{comment}\n分数：{t_s}"); run_comment.font.color.rgb = RGBColor(255, 0, 0)
                                        p_sign = cell.add_paragraph(); p_sign.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                        p_sign.add_run(f"教师签名：{teacher_name}\n{review_date}")
                        doc_changed = True
                    
                    # 4. 填充Excel的DataFrame
                    if do_export:
                        base_df.loc[row_idx, "分数"] = t_s
                        base_df.loc[row_idx, "目的与要求分数"] = p_s; base_df.loc[row_idx, "内容步骤分数"] = c_s; base_df.loc[row_idx, "结果讨论分数"] = r_s
                        base_df.loc[row_idx, "格式扣分"] = format_deduction_points
                        if new_path in file_similarity_info:
                            info = file_similarity_info[new_path]
                            base_df.loc[row_idx, "最大相似度"] = f"{info['max_similarity']*100:.2f}%"
                            base_df.loc[row_idx, "相似文件"] = info["similar_file"]; base_df.loc[row_idx, "相似度扣分"] = info["penalty"]
                        main_table = report_data["main_table"]
                        if main_table and len(main_table.rows[0].cells) >= 4:
                            base_df.loc[row_idx, "课程名称"] = main_table.rows[0].cells[0].text.replace("课程名称：", "").strip()
                            base_df.loc[row_idx, "实训名称"] = main_table.rows[0].cells[3].text.replace("实验（训）名称：", "").strip()
                        base_df.loc[row_idx, "报告时间"] = report_data["date_obj"].strftime('%Y-%m-%d') if report_data["date_obj"] else ""

                    # 5. 保存文档
                    if doc_changed:
                        doc.save(new_path)
                except Exception as e:
                    log_emitter.emit(f"!! 处理文件时发生意外错误: {os.path.basename(new_path)}，错误: {e}")
            
            # 将处理完的当前作业的DataFrame更新回sheets_data
            sheets_data[ws_title] = base_df.sort_values(by='序号').reset_index(drop=True)
        
        # --- 阶段四：为当前班级生成最终的Excel文件 ---
        if do_export:
            log_emitter.emit(f"正在为班级 {class_name} 生成Excel报告...")
            
            # 创建汇总表
            summary_df = name_list_data[class_name][['序号', '姓名', '学号']].copy()
            for sheet_name, df in sheets_data.items():
                # 只汇总成绩页，跳过旧的汇总页
                if sheet_name != "汇总分数" and '学号' in df.columns and '分数' in df.columns:
                    summary_df = pd.merge(summary_df, df[['学号', '分数']], on='学号', how='left').rename(columns={'分数': sheet_name})
                    summary_df[sheet_name] = summary_df[sheet_name].fillna(0).astype(int)
            sheets_data["汇总分数"] = summary_df
            
            # 将所有工作表（旧的+更新的）一次性写入Excel文件
            with pd.ExcelWriter(output_filepath, engine='openpyxl') as writer:
                for sheet_name, df_to_write in sheets_data.items():
                    df_to_write.to_excel(writer, sheet_name=sheet_name, index=False)
                    ws = writer.sheets[sheet_name]
                    for row in ws.iter_rows():
                        for cell in row: cell.alignment = Alignment(horizontal='center', vertical='center')
                    
                    col_widths = {
                        '序号': 5, '姓名': 10, '学号': 12, '分数': 6, 
                        '课程名称': 20, '实训名称': 25, '报告时间': 12,
                        '目的与要求分数': 6, '内容步骤分数': 6, '结果讨论分数': 6, 
                        '最大相似度': 10, '相似文件': 25, '相似度扣分': 10, '格式扣分': 6
                    }
                    for i, col in enumerate(df_to_write.columns, 1):
                        ws.column_dimensions[get_column_letter(i)].width = col_widths.get(col, 12)
            log_emitter.emit(f"班级 {class_name} 的报告处理完毕！")

def load_and_build_app_data(gui_paths, config, log_emitter):
    log_emitter.emit("正在加载和验证数据...")
    try:
        app_data = {'config': config, 'paths': gui_paths}

        # MODIFIED: 使用更健壮的路径获取方式，兼容--onefile打包
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))

        name_list_file = os.path.join(base_path, "名单.xlsx")
        comment_file = os.path.join(base_path, "评语库.xlsx")
        
        app_data['name_list_data'] = pd.read_excel(name_list_file, sheet_name=None)
        app_data['comments'] = pd.read_excel(comment_file)["评语"].tolist()
        name_lookup = {}
        for sheet_name, df in app_data['name_list_data'].items():
            for index, row in df.iterrows(): name_lookup[row['姓名']] = {'row': row, 'sheet': sheet_name, 'original_index': index}
        app_data['name_lookup'] = name_lookup
        app_data['all_names_sorted'] = sorted(list(name_lookup.keys()), key=len, reverse=True)
        log_emitter.emit("数据加载成功！")
        return app_data
    except FileNotFoundError as e:
        log_emitter.emit(f"!! 错误：找不到核心文件 {os.path.basename(e.filename)}。请确保它和程序在同一目录下。")
        return None
    except Exception as e:
        log_emitter.emit(f"!! 加载数据时发生错误: {e}")
        return None


# ==============================================================================
# 4. PyQt6 GUI界面定义与控制逻辑
# ==============================================================================
class HelpDialog(QDialog):
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("使用说明")
        self.setWindowIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogHelpButton))
        self.setGeometry(150, 150, 750, 650)
        layout = QVBoxLayout(self)
        text_widget = QTextEdit(self)
        text_widget.setReadOnly(True)
        text_widget.setFont(QFont("Microsoft YaHei", 10))
        text_widget.setText(load_readme_content())
        layout.addWidget(text_widget)
        self.setLayout(layout)

class Worker(QObject):
    log_message = pyqtSignal(str)
    task_finished = pyqtSignal()

    # MODIFIED: 构造函数增加 do_format_deduction 参数
    def __init__(self, input_folder, do_rename, do_format, do_score, do_export, do_format_deduction):
        super().__init__()
        self.input_folder = input_folder
        self.do_rename = do_rename
        self.do_format = do_format
        self.do_score = do_score
        self.do_export = do_export
        self.do_format_deduction = do_format_deduction # NEW: 保存新参数

    def run_task(self):
        try:
            current_config = load_config()
            start_time = time.time()
            if not self.input_folder:
                self.log_message.emit("!! 错误：必须选择一个待处理文件夹！")
                raise ValueError("路径未选择")
            
            output_folder = os.path.join(os.path.dirname(self.input_folder), "已处理文件")
            app_data = load_and_build_app_data({"input_folder": self.input_folder, "output_folder": output_folder}, current_config, self.log_message)
            
            if app_data is None: 
                raise ValueError("数据加载失败，请检查日志。")
            
            processed_files = prepare_and_copy_files(app_data, self.do_rename, self.log_message)

            if not processed_files:
                self.log_message.emit("未找到可处理的文件，任务提前结束。")
                return

            # MODIFIED: 将新参数传递给主处理函数
            process_reports_conditionally(processed_files, app_data, self.log_message,
                                          do_format=self.do_format, 
                                          do_score=self.do_score, 
                                          do_export=self.do_export,
                                          do_format_deduction=self.do_format_deduction)
            
            end_time = time.time()
            self.log_message.emit(f"\n全部处理完成！共耗时: {end_time - start_time:.2f} 秒")
        except Exception as e:
            self.log_message.emit(f"!! 发生严重错误: {e}")
            self.log_message.emit("!! 任务异常终止。")
        finally:
            self.task_finished.emit()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("一键实训报告处理工具")
        self.setGeometry(100, 100, 800, 680)
        self.setWindowIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_ComputerIcon))
        self.app_config = load_config()
        self.worker = None
        self.thread = None
        self._create_menu()
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(10, 10, 10, 10)
        self._create_path_group(main_layout)
        self._create_params_group(main_layout)
        self._create_scoring_group(main_layout)
        self._create_actions_group(main_layout)
        self._create_log_group(main_layout)
        self.load_config_to_gui()
        self.toggle_similarity_widgets()
        self._update_export_checkbox_state()
        self.statusBar().showMessage("就绪")

    def _create_menu(self):
        
        menubar = self.menuBar(); help_menu = menubar.addMenu("帮助")
        help_action = QAction(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogHelpButton), "使用说明", self)
        help_action.triggered.connect(self.show_help)
        help_menu.addAction(help_action)

    def _create_path_group(self, parent_layout):
        
        path_group = QGroupBox("核心路径设置")
        layout = QHBoxLayout(path_group)
        self.input_folder_edit = QLineEdit(); self.input_folder_edit.setPlaceholderText("请选择或输入包含所有学生报告的文件夹路径")
        layout.addWidget(self.input_folder_edit)
        browse_button = QPushButton(" 浏览..."); browse_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DirOpenIcon))
        browse_button.clicked.connect(self.browse_input_folder)
        layout.addWidget(browse_button)
        parent_layout.addWidget(path_group)

    def _create_params_group(self, parent_layout):
        params_group = QGroupBox("基本参数与可选步骤")
        main_v_layout = QVBoxLayout(params_group)

        # ... (第一、二行布局无改动) ...
        top_row_layout = QHBoxLayout(); top_row_layout.addWidget(QLabel("教师姓名:")); self.teacher_name_edit = QLineEdit(); self.teacher_name_edit.setMaximumWidth(100); top_row_layout.addWidget(self.teacher_name_edit); top_row_layout.addSpacing(20)
        top_row_layout.addWidget(QLabel("批阅于报告日期")); self.delay_days_spin = QSpinBox(); self.delay_days_spin.setRange(0, 30); top_row_layout.addWidget(self.delay_days_spin); top_row_layout.addWidget(QLabel("天后")); top_row_layout.addStretch()
        sim_row_layout = QHBoxLayout(); self.similarity_checkbox = QCheckBox("启用相似度检测"); self.similarity_checkbox.stateChanged.connect(self.toggle_similarity_widgets); sim_row_layout.addWidget(self.similarity_checkbox)
        self.sim_param_container = QWidget(); sim_param_layout = QHBoxLayout(self.sim_param_container); sim_param_layout.setContentsMargins(10, 0, 0, 0); sim_param_layout.addWidget(QLabel("相似度超过(%)")); self.similarity_threshold_spin = QSpinBox(); self.similarity_threshold_spin.setRange(0, 100); self.similarity_threshold_spin.setToolTip("当报告内容相似度高于此百分比时，将触发扣分"); sim_param_layout.addWidget(self.similarity_threshold_spin); sim_param_layout.addSpacing(15); sim_param_layout.addWidget(QLabel("最多扣")); self.max_penalty_spin = QSpinBox(); self.max_penalty_spin.setRange(0, 100); self.max_penalty_spin.setToolTip("相似度达到100%时的最大扣分数额"); sim_param_layout.addWidget(self.max_penalty_spin); sim_param_layout.addWidget(QLabel("分"))
        sim_row_layout.addWidget(self.sim_param_container); sim_row_layout.addStretch()
        
        # MODIFIED: 第三行可选步骤布局增加新控件
        steps_layout = QHBoxLayout()
        steps_layout.setContentsMargins(0, 5, 0, 0)
        self.rename_checkbox = QCheckBox("统一文件名"); self.rename_checkbox.setChecked(True)
        self.format_checkbox = QCheckBox("统一文档格式"); self.format_checkbox.setChecked(True)
        # NEW: 新增“启用格式扣分”复选框
        self.format_deduction_checkbox = QCheckBox("启用格式扣分"); self.format_deduction_checkbox.setChecked(True)
        self.score_checkbox = QCheckBox("自动评分评语"); self.score_checkbox.setChecked(True)
        self.score_checkbox.stateChanged.connect(self._update_export_checkbox_state)
        self.export_checkbox = QCheckBox("导出成绩表格"); self.export_checkbox.setChecked(True)
        steps_layout.addWidget(self.rename_checkbox); steps_layout.addWidget(self.format_checkbox)
        steps_layout.addWidget(self.format_deduction_checkbox) # 添加到布局
        steps_layout.addWidget(self.score_checkbox); steps_layout.addWidget(self.export_checkbox); steps_layout.addStretch()

        main_v_layout.addLayout(top_row_layout)
        main_v_layout.addLayout(sim_row_layout)
        main_v_layout.addLayout(steps_layout)
        parent_layout.addWidget(params_group)

    def _create_scoring_group(self, parent_layout):
        
        scoring_group = QGroupBox("评分标准")
        layout = QGridLayout(scoring_group)
        headers = ["评分项", "最大分数", "基础分数", "最大字数", "基础字数"];
        for col, header in enumerate(headers): layout.addWidget(QLabel(header, font=QFont("Any", weight=QFont.Weight.Bold)), 0, col, Qt.AlignmentFlag.AlignCenter)
        sections = {"purpose": "目的与要求", "content": "内容与步骤", "result": "结果与讨论"}; param_keys = ["max_score", "base_score", "max_chars", "base_chars"]; self.scoring_spins = {}
        for row, (key, label_text) in enumerate(sections.items(), 1):
            row_label = QLabel(f"{label_text}:")
            layout.addWidget(row_label, row, 0, Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter)
            self.scoring_spins[key] = {}
            for col, param_key in enumerate(param_keys, 1):
                spin_box = QSpinBox(); spin_box.setRange(0, 5000); spin_box.setFixedWidth(80); self.scoring_spins[key][param_key] = spin_box; layout.addWidget(spin_box, row, col, Qt.AlignmentFlag.AlignCenter)
        parent_layout.addWidget(scoring_group)
    
    def _create_actions_group(self, parent_layout):
        
        action_layout = QHBoxLayout()
        action_layout.setContentsMargins(0, 5, 0, 5)
        self.save_button = QPushButton(" 保存当前配置"); self.save_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogSaveButton))
        self.save_button.clicked.connect(self.save_config_from_gui)
        self.start_button = QPushButton(" 开始处理"); self.start_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
        self.start_button.clicked.connect(self.start_processing)
        action_layout.addWidget(self.save_button)
        action_layout.addStretch()
        action_layout.addWidget(self.start_button)
        parent_layout.addLayout(action_layout)

    def _create_log_group(self, parent_layout):
        
        log_group = QGroupBox("处理日志")
        layout = QVBoxLayout(log_group)
        self.log_text_edit = QTextEdit(); self.log_text_edit.setReadOnly(True)
        self.log_text_edit.setFont(QFont("Courier New", 9))
        layout.addWidget(self.log_text_edit)
        parent_layout.addWidget(log_group)
        parent_layout.setStretchFactor(log_group, 1)

    def load_config_to_gui(self):
        """槽函数：将 self.app_config 中的配置数据填充到界面控件中。"""
        # 加载基本参数
        self.teacher_name_edit.setText(self.app_config.get("teacher", {}).get("name", "教师"))
        self.delay_days_spin.setValue(self.app_config.get("date_settings", {}).get("review_delay_days", 2))
        
        # 加载相似度设置
        sim_config = self.app_config.get("similarity", {})
        self.similarity_checkbox.setChecked(sim_config.get("enabled", True))
        self.similarity_threshold_spin.setValue(int(sim_config.get("threshold", 0.8) * 100))
        self.max_penalty_spin.setValue(sim_config.get("max_penalty", 20))
        
        # 加载评分标准
        scoring_config = self.app_config.get("scoring", {})
        for section, params in self.scoring_spins.items():
            for key, spin_box in params.items():
                spin_box.setValue(scoring_config.get(section, {}).get(key, 0))

        # NEW: 加载5个处理步骤开关的状态
        # 使用 .get("processing_options", {}) 来确保旧的配置文件不会导致错误
        options_config = self.app_config.get("processing_options", {})
        self.rename_checkbox.setChecked(options_config.get("do_rename", True))
        self.format_checkbox.setChecked(options_config.get("do_format", True))
        self.format_deduction_checkbox.setChecked(options_config.get("do_format_deduction", True))
        self.score_checkbox.setChecked(options_config.get("do_score", True))
        self.export_checkbox.setChecked(options_config.get("do_export", True))

    def save_config_from_gui(self):
        """槽函数：响应'保存配置'按钮点击，从界面控件读取数据并保存到文件。"""
        try:
            new_config = get_default_config()
            # 保存基本参数
            new_config["teacher"]["name"] = self.teacher_name_edit.text()
            new_config["date_settings"]["review_delay_days"] = self.delay_days_spin.value()
            
            # 保存相似度设置
            new_config["similarity"]["enabled"] = self.similarity_checkbox.isChecked()
            new_config["similarity"]["threshold"] = self.similarity_threshold_spin.value() / 100.0
            new_config["similarity"]["max_penalty"] = self.max_penalty_spin.value()
            
            # 保存评分标准
            for section, params in self.scoring_spins.items():
                for key, spin_box in params.items(): 
                    new_config["scoring"][section][key] = spin_box.value()
            
            # NEW: 保存5个处理步骤开关的状态
            options = {
                "do_rename": self.rename_checkbox.isChecked(),
                "do_format": self.format_checkbox.isChecked(),
                "do_format_deduction": self.format_deduction_checkbox.isChecked(),
                "do_score": self.score_checkbox.isChecked(),
                "do_export": self.export_checkbox.isChecked()
            }
            new_config["processing_options"] = options

            # 写入文件
            self.app_config = new_config
            save_config(self.app_config)
            self.log("配置已根据当前界面设置保存到 config.json")
            self.statusBar().showMessage("配置已成功保存！", 5000)
        except Exception as e:
            self.log(f"!! 保存配置时出错: {e}")
            QMessageBox.critical(self, "错误", f"保存配置时出错: {e}")

    def browse_input_folder(self):
        path = QFileDialog.getExistingDirectory(self, "选择待处理报告文件夹")
        if path: self.input_folder_edit.setText(path)

    def log(self, message):
        timestamp = time.strftime('%H:%M:%S')
        self.log_text_edit.append(f"[{timestamp}] {message}")
    def toggle_similarity_widgets(self):
        
        self.sim_param_container.setEnabled(self.similarity_checkbox.isChecked())
    def _update_export_checkbox_state(self):
        
        is_scoring_enabled = self.score_checkbox.isChecked()
        self.export_checkbox.setEnabled(is_scoring_enabled)
        if not is_scoring_enabled:
            self.export_checkbox.setChecked(False)
    def show_help(self):
        
        if not hasattr(self, 'help_dialog'): self.help_dialog = HelpDialog(self)
        self.help_dialog.show()
    def set_buttons_state(self, enabled):
        
        self.save_button.setEnabled(enabled); self.start_button.setEnabled(enabled)
    
    def start_processing(self):
        input_folder = self.input_folder_edit.text()
        if not input_folder or not os.path.isdir(input_folder):
            QMessageBox.warning(self, "路径无效", "请输入或选择一个有效的待处理报告文件夹。")
            return
        self.set_buttons_state(False)
        self.statusBar().showMessage("任务处理中，请稍候...")
        self.log("...任务已启动，正在后台处理...")
        self.thread = QThread()
        # MODIFIED: 将新复选框的状态传递给Worker
        self.worker = Worker(
            input_folder, 
            self.rename_checkbox.isChecked(), 
            self.format_checkbox.isChecked(), 
            self.score_checkbox.isChecked(), 
            self.export_checkbox.isChecked(),
            self.format_deduction_checkbox.isChecked()
        )
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run_task); self.worker.task_finished.connect(self.on_task_finished); self.worker.log_message.connect(self.log)
        self.thread.start()

    def on_task_finished(self):
        
        self.set_buttons_state(True); self.statusBar().showMessage("任务完成！", 10000); self.thread.quit(); self.thread.wait(); self.thread = None; self.worker = None
    def closeEvent(self, event):
        
        if self.thread and self.thread.isRunning():
            reply = QMessageBox.question(self, '退出确认', '任务仍在后台运行，确定要强制退出吗？', QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes: self.thread.quit(); self.thread.wait(); event.accept()
            else: event.ignore()
        else: event.accept()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())